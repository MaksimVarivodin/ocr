from __future__ import annotations

import os
from typing import List, Tuple

from docx import Document as DocxDocument
from docx.shared import Inches

from .model import Document, Page, TextRun, ImageItem
from .buffer import BufferManager


def _extract_images_from_run(run) -> List[bytes]:
    """Return a list of image bytes embedded in this run (if any)."""
    images: List[bytes] = []
    r = run._r
    # Find all blips with embed relationships
    blips = r.xpath('.//a:blip')
    for blip in blips:
        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if not rId:
            continue
        try:
            part = run.part.related_parts[rId]
            images.append(part.blob)
        except Exception:
            continue
    return images


def read_docx(path: str, buffer: BufferManager) -> Document:
    docx = DocxDocument(path)
    page = Page(index=0)
    order = 0
    for para in docx.paragraphs:
        # Walk runs to preserve order of text and inline images
        cur_text_parts: List[str] = []
        for run in para.runs:
            # First, flush any images encountered before/after text chunks
            imgs = _extract_images_from_run(run)
            if imgs:
                # If we have accumulated text, flush as TextRun before the image
                if any(s.strip() for s in cur_text_parts):
                    page.items.append(TextRun(page_index=0, order=order, text=''.join(cur_text_parts).strip()))
                    order += 1
                    cur_text_parts = []
                for idx, blob in enumerate(imgs):
                    # Persist image to buffer with predictable name
                    fname = f"docx-img-{order:04d}-{idx:02d}.png"
                    out_path = buffer.path(fname)
                    try:
                        with open(out_path, 'wb') as f:
                            f.write(blob)
                        page.items.append(ImageItem(page_index=0, order=order, src_path=out_path))
                        order += 1
                    except Exception:
                        # ignore failed image extraction
                        pass
            # Then add run text (if any)
            if run.text:
                cur_text_parts.append(run.text)
        # Flush remaining text for the paragraph
        if any(s.strip() for s in cur_text_parts):
            page.items.append(TextRun(page_index=0, order=order, text=''.join(cur_text_parts).strip()))
            order += 1
    return Document(pages=[page])


def write_docx(doc: Document, out_path: str) -> str:
    d = DocxDocument()
    section = d.sections[0]
    # Available width = page width - (left+right) margins
    avail_width = section.page_width - section.left_margin - section.right_margin
    for item in doc.iter_items():
        if isinstance(item, TextRun):
            p = d.add_paragraph()
            p.add_run(item.translated_text or item.text or "")
        elif isinstance(item, ImageItem):
            img_path = item.translated_path or item.src_path
            try:
                d.add_picture(img_path, width=avail_width)
            except Exception:
                # fallback: put a placeholder paragraph
                p = d.add_paragraph()
                p.add_run(f"[image: {os.path.basename(img_path)}]")
    d.save(out_path)
    return out_path
